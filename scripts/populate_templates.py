#!/usr/bin/env python3
"""
Template population script for EU AI Act compliance documents.

Usage:
    python populate_templates.py risk.json compliance.json --output draft_docs/

What it does:
    1. Loads JSON data from cert-framework commands
    2. Opens Word template files
    3. Replaces {{PLACEHOLDERS}} with real values
    4. Saves populated documents to output directory

What it does NOT do:
    - Fill in [EXPERT INPUT REQUIRED] sections (that's your job)
    - Validate the data (assumes cert-framework output is correct)
    - Handle errors gracefully (if it fails, it fails loudly)
"""

import json
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Error: python-docx not installed")
    print("Install with: pip install python-docx")
    sys.exit(1)


def load_json(filepath):
    """
    Load and parse JSON file.

    Args:
        filepath: Path to JSON file

    Returns:
        dict: Parsed JSON data

    Raises:
        FileNotFoundError: If file doesn't exist
        json.JSONDecodeError: If file isn't valid JSON
    """
    try:
        with open(filepath, encoding="utf-8") as f:
            return json.load(f)
    except FileNotFoundError:
        print(f"Error: File not found: {filepath}")
        raise
    except json.JSONDecodeError as e:
        print(f"Error: Invalid JSON in {filepath}: {e}")
        raise


def replace_placeholder(doc, placeholder, value):
    """
    Replace placeholder text in Word document.

    This searches through ALL paragraphs and ALL table cells,
    looking for the placeholder text and replacing it.

    Args:
        doc: docx.Document object
        placeholder: String to find (e.g., "{{SYSTEM_NAME}}")
        value: String to replace with

    Why this function exists:
        python-docx doesn't have a simple "find and replace" method.
        We have to manually search through paragraph runs (text segments).
        Each paragraph can have multiple runs with different formatting.
        We have to preserve the formatting while replacing text.

    Critical detail:
        We replace text in RUNS, not paragraphs. A paragraph like:
        "System: {{SYSTEM_NAME}}"
        might be split into runs:
        ["System: ", "{{SYSTEM_NAME}}", ""]
        We need to find which run contains our placeholder.
    """
    # Convert value to string, handle None
    value_str = str(value) if value is not None else ""

    # Search in all paragraphs
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Paragraph contains our placeholder, find which run
            for run in paragraph.runs:
                if placeholder in run.text:
                    # Found it - replace only in this run
                    run.text = run.text.replace(placeholder, value_str)

    # Search in all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Each cell contains paragraphs, search those
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, value_str)


def safe_get(data, *keys, default="[Not provided]"):
    """
    Safely navigate nested dictionaries.

    Args:
        data: Dictionary to navigate
        *keys: Sequence of keys to traverse
        default: Value to return if any key doesn't exist

    Returns:
        Value at the nested key, or default if not found

    Example:
        safe_get(data, 'metadata', 'system_name', default='Unknown')

        Instead of:
        data.get('metadata', {}).get('system_name', 'Unknown')

    Why this exists:
        JSON from cert-framework might be missing fields.
        We don't want to crash if a field is missing.
        We want to clearly show "[Not provided]" so you know to fill it in.
    """
    current = data
    for key in keys:
        if isinstance(current, dict):
            current = current.get(key)
            if current is None:
                return default
        else:
            return default
    return current if current is not None else default


def populate_risk_classification(template_path, output_path, risk_data):
    """
    Populate risk classification report template.

    Args:
        template_path: Path to risk_classification_template.docx
        output_path: Where to save populated document
        risk_data: JSON from 'cert classify-system' command

    What this does:
        Opens template, replaces all {{PLACEHOLDERS}} with data from
        risk_data JSON, saves result.

    Data structure expected:
        {
            "classification": {
                "risk_level": "high" | "limited" | "minimal",
                "title": "High-Risk AI System",
                "description": "...",
                "high_risk_indicators": 3
            },
            "matched_categories": [
                {
                    "category": "Employment",
                    "annex_reference": "Annex III (4)(a)"
                }
            ],
            "requirements": ["req1", "req2"]
        }
    """
    print("Populating risk classification report...")

    # Load template
    doc = Document(template_path)

    # Extract data from JSON
    classification = risk_data.get("classification", {})
    matched_categories = risk_data.get("matched_categories", [])
    requirements = risk_data.get("requirements", [])

    # Replace basic placeholders
    replace_placeholder(doc, "{{GENERATION_DATE}}", datetime.now().strftime("%Y-%m-%d"))
    replace_placeholder(doc, "{{SYSTEM_NAME}}", safe_get(risk_data, "metadata", "system_name"))
    replace_placeholder(
        doc, "{{SYSTEM_VERSION}}", safe_get(risk_data, "metadata", "system_version", default="v1.0")
    )
    replace_placeholder(doc, "{{PROVIDER_NAME}}", safe_get(risk_data, "metadata", "provider_name"))

    # Classification results
    replace_placeholder(
        doc, "{{RISK_LEVEL}}", classification.get("risk_level", "unknown").upper()
    )
    replace_placeholder(
        doc, "{{RISK_INDICATORS}}", str(classification.get("high_risk_indicators", 0))
    )
    replace_placeholder(doc, "{{RISK_TITLE}}", classification.get("title", ""))
    replace_placeholder(doc, "{{RISK_DESCRIPTION}}", classification.get("description", ""))

    # Format matched categories as bullet list
    if matched_categories:
        categories_text = "\n".join(
            [f"• {cat['category']}: {cat['annex_reference']}" for cat in matched_categories]
        )
    else:
        categories_text = "No high-risk categories matched (minimal risk system)"
    replace_placeholder(doc, "{{MATCHED_CATEGORIES}}", categories_text)

    # Format requirements as bullet list
    if requirements:
        requirements_text = "\n".join([f"• {req}" for req in requirements])
    else:
        requirements_text = "• Basic transparency obligations (Article 13, 50)"
    replace_placeholder(doc, "{{REQUIREMENTS_LIST}}", requirements_text)

    # Save populated document
    doc.save(output_path)
    print(f"✓ Saved to {output_path}")


def populate_annex_iv(template_path, output_path, compliance_data, risk_data):
    """
    Populate Annex IV technical documentation template.

    Args:
        template_path: Path to annex_iv_template.docx
        output_path: Where to save populated document
        compliance_data: JSON from 'cert generate-docs' command
        risk_data: JSON from 'cert classify-system' command

    This is the big one - populates 20-25 page document with:
    - System metadata (name, version, provider)
    - Risk classification results
    - Performance metrics from trace analysis
    - Architecture information (if provided)
    - Data governance information (if provided)

    Many sections will still have [EXPERT INPUT REQUIRED] because
    they need human judgment, not automated data.
    """
    print("Populating Annex IV documentation...")

    doc = Document(template_path)

    # Extract data structures
    metadata = compliance_data.get("metadata", {})
    article_15 = compliance_data.get("article_15_compliance", {})
    metrics = article_15.get("metrics", {})
    annex_iv = compliance_data.get("annex_iv_documentation", {})
    sections = annex_iv.get("sections", {})
    trace_summary = compliance_data.get("trace_summary", {})

    # Basic metadata (appears throughout document)
    replace_placeholder(doc, "{{GENERATION_DATE}}", datetime.now().strftime("%Y-%m-%d"))
    replace_placeholder(doc, "{{SYSTEM_NAME}}", metadata.get("system_name", "[Not provided]"))
    replace_placeholder(doc, "{{SYSTEM_VERSION}}", metadata.get("system_version", "v1.0"))
    replace_placeholder(doc, "{{PROVIDER_NAME}}", metadata.get("provider_name", "[Not provided]"))

    # Section 1: General Information
    section1_data = sections.get("section_1_general", {}).get("data", {})
    replace_placeholder(
        doc, "{{INTENDED_PURPOSE}}", section1_data.get("intended_purpose", "[Not provided]")
    )
    replace_placeholder(
        doc, "{{DEPLOYMENT_DATE}}", section1_data.get("deployment_date", "[Not provided]")
    )

    # Risk classification (from risk_data)
    classification = risk_data.get("classification", {})
    risk_level = classification.get("risk_level", "unknown").upper()
    replace_placeholder(doc, "{{RISK_LEVEL}}", risk_level)
    replace_placeholder(doc, "{{RISK_TITLE}}", classification.get("title", ""))

    # Section 2: Architecture
    section2_data = sections.get("section_2_architecture", {}).get("data", {})
    replace_placeholder(doc, "{{MODEL_TYPE}}", section2_data.get("model_type", "[Not provided]"))
    replace_placeholder(
        doc, "{{MODEL_VERSION}}", section2_data.get("model_version", "[Not provided]")
    )
    replace_placeholder(
        doc, "{{INFRASTRUCTURE_DESCRIPTION}}", section2_data.get("infrastructure", "[Not provided]")
    )

    # Section 3: Data Governance
    section3_data = sections.get("section_3_data_governance", {}).get("data", {})
    replace_placeholder(
        doc, "{{TRAINING_DATA_DESCRIPTION}}", section3_data.get("training_data", "[Not provided]")
    )

    # Section 4: Performance Metrics (Article 15)
    replace_placeholder(doc, "{{TOTAL_TRACES}}", str(trace_summary.get("total_traces", 0)))
    replace_placeholder(
        doc,
        "{{TRACE_START_DATE}}",
        trace_summary.get("date_range", {}).get("start", "[Not provided]"),
    )
    replace_placeholder(
        doc,
        "{{TRACE_END_DATE}}",
        trace_summary.get("date_range", {}).get("end", "[Not provided]"),
    )

    replace_placeholder(doc, "{{TOTAL_REQUESTS}}", str(metrics.get("total_requests", 0)))
    replace_placeholder(
        doc, "{{SUCCESSFUL_REQUESTS}}", str(metrics.get("successful_requests", 0))
    )
    replace_placeholder(doc, "{{FAILED_REQUESTS}}", str(metrics.get("failed_requests", 0)))

    # Format error rate as percentage
    error_rate = metrics.get("error_rate", 0.0)
    replace_placeholder(doc, "{{ERROR_RATE}}", f"{error_rate * 100:.2f}")

    # Format response time
    avg_time = metrics.get("avg_response_time_ms", 0.0)
    replace_placeholder(doc, "{{AVG_RESPONSE_TIME}}", f"{avg_time:.1f}")

    # Accuracy metrics (these might not be calculated yet)
    # If not available, placeholders remain as hints to fill in manually
    replace_placeholder(doc, "{{MEAN_CONFIDENCE}}", "[To be calculated from trace analysis]")
    replace_placeholder(doc, "{{ACCURACY_PCT}}", "[To be calculated from trace analysis]")
    replace_placeholder(doc, "{{ROC_AUC}}", "[To be calculated if applicable]")

    # Section 8: Version History
    version_history = (
        f"{metadata.get('system_version', 'v1.0')} "
        f"({datetime.now().strftime('%Y-%m-%d')}): Initial deployment"
    )
    replace_placeholder(doc, "{{VERSION_HISTORY}}", version_history)

    doc.save(output_path)
    print(f"✓ Saved to {output_path}")


def copy_template_as_is(template_path, output_path, doc_name):
    """
    Copy a template without modification.

    Use this for templates that don't have automated placeholders,
    just [EXPERT INPUT REQUIRED] sections.

    Args:
        template_path: Source template
        output_path: Destination
        doc_name: Human-readable name for logging
    """
    import shutil

    shutil.copy(template_path, output_path)
    print(f"✓ Copied {doc_name} (requires manual population)")


def main():
    """
    Main execution function.

    Parses command line arguments, loads data, populates templates.
    """
    # Parse arguments
    if len(sys.argv) < 4:
        print(
            "Usage: python populate_templates.py <risk_json> <compliance_json> --output <output_dir>"
        )
        print("\nExample:")
        print(
            "  python populate_templates.py client_risk.json client_compliance.json --output draft_docs/"
        )
        print("\nThis script:")
        print("  1. Loads JSON from cert-framework commands")
        print("  2. Opens Word templates")
        print("  3. Replaces {{PLACEHOLDERS}} with real data")
        print("  4. Saves populated documents for expert review")
        sys.exit(1)

    risk_json_path = sys.argv[1]
    compliance_json_path = sys.argv[2]

    if "--output" not in sys.argv:
        print("Error: --output directory required")
        sys.exit(1)

    output_dir = Path(sys.argv[sys.argv.index("--output") + 1])
    output_dir.mkdir(parents=True, exist_ok=True)

    # Load data files
    print("Loading data files...")
    print("=" * 70)

    try:
        risk_data = load_json(risk_json_path)
        print("✓ Loaded risk classification data")

        compliance_data = load_json(compliance_json_path)
        print("✓ Loaded compliance data")

    except (FileNotFoundError, json.JSONDecodeError) as e:
        print(f"\nFailed to load data: {e}")
        sys.exit(1)

    # Find templates directory
    # Assumes script is in cert-framework/scripts/
    # Templates are in cert-framework/templates/
    script_dir = Path(__file__).parent
    templates_dir = script_dir.parent / "templates"

    if not templates_dir.exists():
        print(f"\nError: Templates directory not found: {templates_dir}")
        print("Make sure templates/ directory exists with .docx files")
        sys.exit(1)

    # Populate each template
    print("\nPopulating templates...")
    print("=" * 70)

    # 1. Risk Classification Report
    try:
        populate_risk_classification(
            templates_dir / "risk_classification_template.docx",
            output_dir / "Risk_Classification_Report.docx",
            risk_data,
        )
    except Exception as e:
        print(f"✗ Failed to populate risk classification: {e}")

    # 2. Annex IV Technical Documentation
    try:
        populate_annex_iv(
            templates_dir / "annex_iv_template.docx",
            output_dir / "Annex_IV_Technical_Documentation.docx",
            compliance_data,
            risk_data,
        )
    except Exception as e:
        print(f"✗ Failed to populate Annex IV: {e}")

    # 3-5. Copy remaining templates (they don't have automated placeholders)
    remaining_templates = [
        (
            "audit_trail_template.docx",
            "Audit_Trail_Setup_Guide.docx",
            "Audit Trail Guide",
        ),
        (
            "monitoring_framework_template.docx",
            "Monitoring_Framework.docx",
            "Monitoring Framework",
        ),
        (
            "conformity_checklist_template.docx",
            "Conformity_Assessment_Checklist.docx",
            "Conformity Checklist",
        ),
    ]

    for template_name, output_name, display_name in remaining_templates:
        template_path = templates_dir / template_name
        if template_path.exists():
            try:
                copy_template_as_is(template_path, output_dir / output_name, display_name)
            except Exception as e:
                print(f"✗ Failed to copy {display_name}: {e}")
        else:
            print(f"⚠ Template not found: {template_name} (skipping)")

    print("=" * 70)
    print(f"\n✓ All documents generated in {output_dir}")
    print("\nNext steps:")
    print("1. Open each .docx file in Word or LibreOffice")
    print("2. Search for '[EXPERT INPUT REQUIRED]' (Ctrl+F)")
    print("3. Fill in expert commentary and assessments")
    print("4. Review all auto-populated metrics for accuracy")
    print("5. Save final versions and export to PDF")
    print("\n⏱  Estimated expert review time: 8-10 hours")


if __name__ == "__main__":
    main()
